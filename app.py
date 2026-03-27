import os
import io
import re
import base64
import json
import tempfile
import shutil
from collections import Counter
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from flask import Flask, render_template, request, jsonify, Response, send_file
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sequences import PRP_SEQUENCES, GROUPS_ORDER
from blast_runner import (start_job, get_job_queue, cleanup_job,
                          get_kept_files, consume_kept_files, delete_kept_files,
                          cancel_job, _kept_files as kept_files_registry)

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024 * 1024  # 10 GB


# ── Kaplan-Meier helpers ───────────────────────────────────────────

def parse_numbers(text):
    parts = re.split(r'[,;\s\t\n\r]+', text.strip())
    numbers = []
    for part in parts:
        part = part.strip()
        if part:
            try:
                numbers.append(float(part))
            except ValueError:
                raise ValueError(f"'{part}' no es un número válido")
    return numbers


def compute_stats(data, table_rows):
    n = len(data)
    mean_val = sum(data) / n
    if n > 1:
        variance = sum((x - mean_val) ** 2 for x in data) / (n - 1)
        std_val = variance ** 0.5
        sem_val = std_val / (n ** 0.5)
    else:
        std_val = 0.0
        sem_val = 0.0
    km_median = None
    for row in table_rows:
        if row['survival'] <= 0.5:
            km_median = row['time']
            break
    return {'n': n, 'mean': round(mean_val, 4), 'std': round(std_val, 4),
            'sem': round(sem_val, 4), 'km_median': km_median}


def process_survival_data(data):
    counter = Counter(data)
    sorted_data = sorted(counter.items(), key=lambda x: x[0])
    sorted_data = [(count, value) for value, count in counter.items()]
    sorted_data.sort(key=lambda x: x[1])

    total = sum(c for c, _ in sorted_data)
    at_risk = total
    surv = 1.0
    times = [0]
    survivals = [1.0]
    table_rows = []

    for deaths, time in sorted_data:
        surv *= (at_risk - deaths) / at_risk
        times.extend([time, time])
        survivals.extend([survivals[-1], surv])
        table_rows.append({'deaths': int(deaths), 'time': time,
                           'at_risk': int(at_risk), 'survival': round(surv, 4)})
        at_risk -= deaths

    stats = compute_stats(data, table_rows)

    plt.figure(figsize=(10, 6))
    plt.plot(times, survivals, 'b-', linewidth=2)
    plt.scatter(times[1::2], survivals[1::2], color='blue', zorder=5)
    plt.grid(True, linestyle='--', alpha=0.7)
    plt.xlabel('Tiempo (días)', fontsize=12)
    plt.ylabel('Probabilidad de supervivencia', fontsize=12)
    plt.title('Curva de Kaplan-Meier', fontsize=14)
    plt.ylim(-0.05, 1.05)

    img_bytes = io.BytesIO()
    plt.savefig(img_bytes, format='png', dpi=300, bbox_inches='tight')
    plt.close()
    img_bytes.seek(0)

    km_median_str = str(stats['km_median']) if stats['km_median'] is not None else 'N/D'
    text_output = (
        f"# Estadísticas descriptivas\nN\t{stats['n']}\nMedia\t{stats['mean']:.4f}\n"
        f"Desv. típica\t{stats['std']:.4f}\nSEM\t{stats['sem']:.4f}\n"
        f"Mediana KM\t{km_median_str}\n\n# Tabla Kaplan-Meier\n"
        f"Muertes\tTiempo\tAnimales en riesgo\tProbabilidad de supervivencia\n"
    )
    for row in table_rows:
        text_output += f"{row['deaths']}\t{row['time']}\t{row['at_risk']}\t{row['survival']:.4f}\n"

    return img_bytes, text_output, table_rows, stats


# ── Rutas ─────────────────────────────────────────────────────────

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/kaplan-meier', methods=['GET', 'POST'])
def kaplan_meier():
    if request.method == 'GET':
        return render_template('kaplan_meier.html')
    try:
        numbers_text = ''
        if 'file' in request.files and request.files['file'].filename:
            numbers_text = request.files['file'].read().decode('utf-8')
        else:
            numbers_text = request.form.get('numbers', '')
        if not numbers_text.strip():
            return jsonify({'error': 'No se han introducido datos.'}), 400
        numbers = parse_numbers(numbers_text)
        if not numbers:
            return jsonify({'error': 'No se encontraron números válidos.'}), 400
        img_bytes, text_output, table_rows, stats = process_survival_data(numbers)
        img_b64 = base64.b64encode(img_bytes.getvalue()).decode('utf-8')
        return jsonify({'image': img_b64, 'text_output': text_output,
                        'table': table_rows, 'stats': stats})
    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        return jsonify({'error': f'Error inesperado: {str(e)}'}), 500


# ── PRNP-OrthoMiner ───────────────────────────────────────────────

@app.route('/prnp-orthominer')
def prnp_orthominer():
    sequences_json = json.dumps(PRP_SEQUENCES, ensure_ascii=False)
    return render_template('prnp_orthominer.html',
                           sequences_json=sequences_json,
                           groups_order=GROUPS_ORDER)


@app.route('/prnp-orthominer/start', methods=['POST'])
def prnp_start():
    # ── Modo subida de archivos (multipart/form-data) ──
    if request.files:
        species      = request.form.get('nombreEspecie', '').strip()
        ref_name     = request.form.get('refName', '').strip()
        ref_sequence = request.form.get('refSequence', '').strip()

        if not all([species, ref_name, ref_sequence]):
            return jsonify({'error': 'Faltan parámetros obligatorios.'}), 400

        gz_files = [f for f in request.files.getlist('genomeFiles')
                    if f.filename.endswith('.gz')]
        if not gz_files:
            return jsonify({'error': 'No se encontraron archivos .gz entre los subidos.'}), 400

        work_dir    = tempfile.mkdtemp(prefix='prnp_')
        species_dir = os.path.join(work_dir, species)
        os.makedirs(species_dir, exist_ok=True)

        for f in gz_files:
            filename = os.path.basename(f.filename)
            f.save(os.path.join(species_dir, filename))

        job_id = start_job(work_dir, species, ref_name, ref_sequence,
                           work_dir, cleanup_dir=work_dir)
        return jsonify({'job_id': job_id})

    # ── Modo ruta local (JSON) — para uso local ──
    data         = request.get_json(force=True) or {}
    base_path    = data.get('rutaBase', '').strip()
    species      = data.get('nombreEspecie', '').strip()
    ref_name     = data.get('refName', '').strip()
    ref_sequence = data.get('refSequence', '').strip()
    output_path  = data.get('rutaSalida', base_path).strip() or base_path

    if not all([base_path, species, ref_name, ref_sequence]):
        return jsonify({'error': 'Faltan parámetros obligatorios.'}), 400

    job_id = start_job(base_path, species, ref_name, ref_sequence, output_path)
    return jsonify({'job_id': job_id})


@app.route('/prnp-orthominer/stream/<job_id>')
def prnp_stream(job_id):
    q = get_job_queue(job_id)
    if q is None:
        return 'Job not found', 404

    def generate():
        try:
            while True:
                try:
                    msg = q.get(timeout=15)   # espera max 15 s
                except Exception:
                    # Ping real (no comentario) para que onmessage lo reciba y resetee timers
                    yield f"data: {json.dumps({'type': 'ping'})}\n\n"
                    continue
                if msg is None:               # sentinel: pipeline terminado
                    yield f"data: {json.dumps({'type': 'closed'})}\n\n"
                    break
                yield f"data: {json.dumps(msg, ensure_ascii=False)}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"
        finally:
            cleanup_job(job_id)

    return Response(
        generate(),
        mimetype='text/event-stream',
        headers={
            'Cache-Control':    'no-cache',
            'X-Accel-Buffering': 'no',
            'Connection':       'keep-alive'
        }
    )


@app.route('/prnp-orthominer/retry/<job_id>', methods=['POST'])
def prnp_retry(job_id):
    """Reintenta el análisis sobre los archivos conservados, con otra referencia."""
    data         = request.get_json(force=True) or {}
    ref_name     = data.get('refName', '').strip()
    ref_sequence = data.get('refSequence', '').strip()

    if not all([ref_name, ref_sequence]):
        return jsonify({'error': 'Faltan refName o refSequence.'}), 400

    stored = consume_kept_files(job_id)   # quita la entrada sin borrar archivos
    if not stored:
        return jsonify({'error': 'Archivos no disponibles (ya borrados o job no encontrado).'}), 404

    work_dir = stored['work_dir']
    species  = stored['species']

    new_job_id = start_job(work_dir, species, ref_name, ref_sequence,
                           work_dir, cleanup_dir=work_dir)
    return jsonify({'job_id': new_job_id, 'species': species})


@app.route('/prnp-orthominer/delete-files/<job_id>', methods=['POST'])
def prnp_delete_files(job_id):
    """Borra manualmente los archivos temporales conservados tras un análisis sin hit."""
    deleted = delete_kept_files(job_id)
    return jsonify({'deleted': deleted})


@app.route('/prnp-orthominer/cancel/<job_id>', methods=['POST'])
def prnp_cancel(job_id):
    """Señala al pipeline que cancele el análisis en curso."""
    cancel_job(job_id)
    return jsonify({'cancelled': True})


@app.route('/prnp-orthominer/server-status')
def prnp_server_status():
    """Devuelve cuántos directorios temporales están retenidos en el servidor."""
    entries = []
    total_bytes = 0
    for job_id, info in kept_files_registry.items():
        work_dir = info.get('work_dir', '')
        size = 0
        if os.path.isdir(work_dir):
            for dirpath, _, filenames in os.walk(work_dir):
                for fn in filenames:
                    try:
                        size += os.path.getsize(os.path.join(dirpath, fn))
                    except OSError:
                        pass
        total_bytes += size
        entries.append({'job_id': job_id, 'species': info.get('species', '?'),
                        'size_mb': round(size / 1024 / 1024, 1)})
    return jsonify({'count': len(entries), 'total_mb': round(total_bytes / 1024 / 1024, 1),
                    'entries': entries})


@app.route('/prnp-orthominer/cleanup-all', methods=['POST'])
def prnp_cleanup_all():
    """Borra todos los directorios temporales retenidos."""
    deleted = 0
    for job_id in list(kept_files_registry.keys()):
        if delete_kept_files(job_id):
            deleted += 1
    return jsonify({'deleted': deleted})


@app.route('/prnp-orthominer/export-docx', methods=['POST'])
def prnp_export_docx():
    """Genera un documento Word con los resultados del análisis."""
    data = request.get_json()
    species   = data.get('species', 'Especie desconocida')
    taxonomy  = data.get('taxonomy') or {}
    results   = data.get('results', [])

    doc = Document()

    # Título
    title = doc.add_heading(species, level=1)
    title.runs[0].font.color.rgb = RGBColor(0x6b, 0x46, 0xc1)

    # Taxonomía
    if taxonomy.get('genero'):
        doc.add_heading('Clasificación taxonómica', level=2)
        tbl = doc.add_table(rows=3, cols=2)
        tbl.style = 'Table Grid'
        for i, (label, key) in enumerate([('Género', 'genero'), ('Familia', 'familia'), ('Orden', 'orden')]):
            tbl.rows[i].cells[0].text = label
            tbl.rows[i].cells[1].text = taxonomy.get(key, '—')

    for idx, res in enumerate(results, 1):
        doc.add_heading(f'Resultado {idx}' + (f' — {res.get("genome_file","")}' if res.get('genome_file') else ''), level=2)

        # Coordenadas
        if res.get('coords_summary'):
            doc.add_heading('Coordenadas genómicas', level=3)
            p = doc.add_paragraph()
            run = p.add_run(res['coords_summary'])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

        # Información genómica
        if res.get('genomic_info'):
            doc.add_heading('Información genómica', level=3)
            p = doc.add_paragraph()
            run = p.add_run(res['genomic_info'])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

        # Secuencia directa
        if res.get('fasta_direct'):
            doc.add_heading('Secuencia extendida (directa)', level=3)
            p = doc.add_paragraph()
            run = p.add_run(res['fasta_direct'])
            run.font.name = 'Courier New'
            run.font.size = Pt(8)

        # Secuencia RC
        if res.get('fasta_rc'):
            doc.add_heading('Secuencia extendida (reversa complementaria)', level=3)
            p = doc.add_paragraph()
            run = p.add_run(res['fasta_rc'])
            run.font.name = 'Courier New'
            run.font.size = Pt(8)

        if idx < len(results):
            doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    safe_name = re.sub(r'[^\w\s-]', '', species).strip().replace(' ', '_') or 'resultado'
    return send_file(buf, as_attachment=True,
                     download_name=f'PRNP_{safe_name}.docx',
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
